from io import BytesIO
from zipfile import BadZipFile, ZipFile

import fitz
import pdfplumber
from docx import Document

MAX_DOCUMENT_PAGES = 50
MAX_EXTRACTED_CHARACTERS = 100_000
MAX_DOCX_ARCHIVE_ENTRIES = 1_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MIN_USEFUL_TEXT_CHARACTERS = 40


class ResumeParsingError(ValueError):
    """Raised when a resume cannot be safely converted into useful text."""


class ResumeParser:
    def extract_text(self, content: bytes, mime_type: str) -> str:
        if mime_type == "application/pdf":
            text = self._extract_pdf(content)
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            text = self._extract_docx(content)
        elif mime_type == "application/msword":
            raise ResumeParsingError(
                "Legacy DOC files require conversion to PDF or DOCX before parsing."
            )
        else:
            raise ResumeParsingError("Unsupported resume file type.")

        normalized = self._normalize_text(text)
        if len(normalized) < MIN_USEFUL_TEXT_CHARACTERS:
            raise ResumeParsingError(
                "The resume contains too little extractable text. "
                "It may be scanned or image-only."
            )
        return normalized[:MAX_EXTRACTED_CHARACTERS]

    def _extract_pdf(self, content: bytes) -> str:
        try:
            with fitz.open(stream=content, filetype="pdf") as document:
                if document.needs_pass:
                    raise ResumeParsingError("Password-protected PDFs are not supported.")
                if document.page_count > MAX_DOCUMENT_PAGES:
                    raise ResumeParsingError(
                        f"PDF resumes cannot exceed {MAX_DOCUMENT_PAGES} pages."
                    )
                text = "\n".join(page.get_text("text") for page in document)
        except ResumeParsingError:
            raise
        except Exception:
            text = ""

        if len(self._normalize_text(text)) >= MIN_USEFUL_TEXT_CHARACTERS:
            return text

        try:
            with pdfplumber.open(BytesIO(content)) as document:
                if len(document.pages) > MAX_DOCUMENT_PAGES:
                    raise ResumeParsingError(
                        f"PDF resumes cannot exceed {MAX_DOCUMENT_PAGES} pages."
                    )
                return "\n".join(page.extract_text() or "" for page in document.pages)
        except ResumeParsingError:
            raise
        except Exception as exception:
            raise ResumeParsingError("The PDF resume could not be read.") from exception

    def _extract_docx(self, content: bytes) -> str:
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                    raise ResumeParsingError("The DOCX archive contains too many files.")
                if (
                    sum(entry.file_size for entry in entries)
                    > MAX_DOCX_UNCOMPRESSED_BYTES
                ):
                    raise ResumeParsingError(
                        "The uncompressed DOCX document exceeds the safety limit."
                    )
            document = Document(BytesIO(content))
        except ResumeParsingError:
            raise
        except (BadZipFile, ValueError, KeyError) as exception:
            raise ResumeParsingError("The DOCX resume could not be read.") from exception

        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [" ".join(line.split()) for line in text.replace("\x00", "").splitlines()]
        return "\n".join(line for line in lines if line).strip()
